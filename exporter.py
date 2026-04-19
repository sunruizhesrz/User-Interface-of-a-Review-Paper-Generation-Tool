"""
exporter.py - 多格式导出模块（Sprint 2 完整版）
成员 F：综述论文生成工具系统集成

Sprint 2 升级内容（F-03 实现完整多格式导出）：
  - Markdown：直接写入（Sprint 1 已完成，保持）
  - PDF：pandoc 优先 → reportlab fallback（完整实现，支持中文）
  - Word：python-docx 完整实现（含粗体、斜体、引用、标题层级）
"""

import os
import re
import subprocess
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger("exporter")


# ============================================================
# Markdown 导出（Sprint 1 已完成，保持不变）
# ============================================================

def export_markdown(content: str, output_path: str = "survey.md") -> str:
    """将综述内容保存为 Markdown 文件"""
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    logger.info(f"Markdown 导出成功：{path.resolve()}")
    return str(path.resolve())


# ============================================================
# PDF 导出（Sprint 2 完整实现）
# ============================================================

def _check_pandoc() -> bool:
    """检查 pandoc 是否可用"""
    try:
        result = subprocess.run(["pandoc", "--version"], capture_output=True, timeout=5)
        return result.returncode == 0
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


def _export_pdf_pandoc(md_path: str, output_path: str) -> str:
    """使用 pandoc 将 Markdown 转换为 PDF"""
    cmd = [
        "pandoc", str(md_path), "-o", str(output_path),
        "--pdf-engine=xelatex",
        "-V", "CJKmainfont=Microsoft YaHei",     # 中文字体（Windows）
        "-V", "geometry:margin=2.5cm",
        "-V", "fontsize=12pt",
        "--toc",                                   # 自动生成目录
    ]
    result = subprocess.run(cmd, capture_output=True, timeout=120)
    if result.returncode != 0:
        raise RuntimeError(f"pandoc 失败：{result.stderr.decode('utf-8', errors='replace')}")
    return str(output_path)


def _export_pdf_reportlab(content: str, output_path: str) -> str:
    """使用 reportlab 生成 PDF（中文支持版）"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # 尝试注册中文字体
    _register_cn_font()

    path = Path(output_path)
    doc = SimpleDocTemplate(
        str(path), pagesize=A4,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm
    )

    styles = getSampleStyleSheet()
    cn_font = _get_cn_font_name()

    title_style   = ParagraphStyle("SurveyTitle",   parent=styles["Title"],   fontName=cn_font, fontSize=18, spaceAfter=12, alignment=TA_CENTER)
    h1_style      = ParagraphStyle("SurveyH1",      parent=styles["Heading1"], fontName=cn_font, fontSize=14, spaceBefore=16, spaceAfter=8)
    h2_style      = ParagraphStyle("SurveyH2",      parent=styles["Heading2"], fontName=cn_font, fontSize=12, spaceBefore=12, spaceAfter=6)
    h3_style      = ParagraphStyle("SurveyH3",      parent=styles["Heading3"], fontName=cn_font, fontSize=11, spaceBefore=8,  spaceAfter=4)
    body_style    = ParagraphStyle("SurveyBody",    parent=styles["Normal"],   fontName=cn_font, fontSize=10.5, leading=18, spaceAfter=4)
    ref_style     = ParagraphStyle("SurveyRef",     parent=styles["Normal"],   fontName=cn_font, fontSize=9.5, leading=15, leftIndent=12)

    story = []
    in_ref_section = False

    for line in content.split("\n"):
        raw = line.rstrip()

        if raw.startswith("# "):
            story.append(Paragraph(_md_to_xml(raw[2:]), title_style))
            story.append(Spacer(1, 0.3*cm))
        elif raw.startswith("## "):
            story.append(Spacer(1, 0.2*cm))
            story.append(Paragraph(_md_to_xml(raw[3:]), h1_style))
            if "参考文献" in raw or "References" in raw:
                in_ref_section = True
        elif raw.startswith("### "):
            story.append(Paragraph(_md_to_xml(raw[4:]), h2_style))
        elif raw.startswith("#### "):
            story.append(Paragraph(_md_to_xml(raw[5:]), h3_style))
        elif raw.startswith("---"):
            story.append(HRFlowable(width="100%", thickness=0.5, color="#cccccc", spaceAfter=6))
        elif raw.startswith("- ") or raw.startswith("* "):
            bullet = "• " + raw[2:]
            story.append(Paragraph(_md_to_xml(bullet), body_style))
        elif re.match(r"^\[?\d+\]", raw) and in_ref_section:
            story.append(Paragraph(_md_to_xml(raw), ref_style))
        elif raw.strip():
            story.append(Paragraph(_md_to_xml(raw), body_style))
        else:
            story.append(Spacer(1, 0.2*cm))

    doc.build(story)
    logger.info(f"PDF 导出成功（reportlab）：{path.resolve()}")
    return str(path.resolve())


def _md_to_xml(text: str) -> str:
    """将 Markdown 内联标记转换为 ReportLab XML 标签（防注入）"""
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.+?)\*",     r"<i>\1</i>", text)
    text = re.sub(r"`(.+?)`",       r"<font name='Courier'>\1</font>", text)
    return text


def _register_cn_font():
    """尝试注册 Windows/macOS 中文字体"""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = [
        ("MSYH", "C:/Windows/Fonts/msyh.ttc"),
        ("SimSun", "C:/Windows/Fonts/simsun.ttc"),
        ("PingFang", "/System/Library/Fonts/PingFang.ttc"),
        ("Helvetica", None),   # 最终降级
    ]
    for name, path in candidates:
        if path and Path(path).exists():
            try:
                pdfmetrics.registerFont(TTFont(name, path))
                _register_cn_font._name = name
                return
            except Exception:
                continue
    _register_cn_font._name = "Helvetica"


_register_cn_font._name = "Helvetica"


def _get_cn_font_name() -> str:
    _register_cn_font()
    return _register_cn_font._name


def export_pdf(content: str, output_path: str = "survey.pdf") -> str:
    """
    PDF 导出主入口：pandoc → reportlab fallback

    Args:
        content    : Markdown 格式综述文本
        output_path: 输出文件路径
    Returns:
        str: 输出文件绝对路径
    """
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    # 先将 Markdown 写到临时文件
    tmp_md = path.with_suffix(".tmp.md")
    tmp_md.write_text(content, encoding="utf-8")

    try:
        if _check_pandoc():
            result = _export_pdf_pandoc(str(tmp_md), str(path))
            tmp_md.unlink(missing_ok=True)
            return result
        else:
            logger.info("pandoc 不可用，改用 reportlab")
    except Exception as e:
        logger.warning(f"pandoc 失败：{e}，改用 reportlab")

    # reportlab 方案
    try:
        result = _export_pdf_reportlab(content, str(path))
        tmp_md.unlink(missing_ok=True)
        return result
    except ImportError:
        tmp_md.unlink(missing_ok=True)
        raise RuntimeError("PDF 导出失败：pandoc 和 reportlab 均不可用。请运行 pip install reportlab")
    except Exception as e:
        tmp_md.unlink(missing_ok=True)
        raise RuntimeError(f"PDF 导出失败：{e}")


# ============================================================
# Word 导出（Sprint 2 完整实现）
# ============================================================

def export_docx(content: str, output_path: str = "survey.docx") -> str:
    """
    将 Markdown 综述内容转换为 Word (.docx) 文件。
    支持：标题层级（H1-H4）、粗体、斜体、列表、分隔线、引用格式。

    Args:
        content    : Markdown 格式综述文本
        output_path: 输出文件路径
    Returns:
        str: 输出文件绝对路径
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        raise RuntimeError("python-docx 未安装，请运行：pip install python-docx")

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ── 页面设置 ─────────────────────────────────────────
    section = doc.sections[0]
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── 默认字体 ─────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    # 中文字体兜底
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "宋体")

    def _add_cn_font(run):
        """为 run 设置中文字体"""
        run.font.name = "Times New Roman"
        run._element.rPr.get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")

    in_ref_section = False

    for line in content.split("\n"):
        raw = line.rstrip()

        if raw.startswith("# "):
            p = doc.add_heading(raw[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                _add_cn_font(run)
            in_ref_section = False

        elif raw.startswith("## "):
            p = doc.add_heading(raw[3:], level=1)
            for run in p.runs:
                _add_cn_font(run)
            if "参考文献" in raw or "References" in raw:
                in_ref_section = True

        elif raw.startswith("### "):
            p = doc.add_heading(raw[4:], level=2)
            for run in p.runs:
                _add_cn_font(run)

        elif raw.startswith("#### "):
            p = doc.add_heading(raw[5:], level=3)
            for run in p.runs:
                _add_cn_font(run)

        elif raw.startswith("---"):
            doc.add_paragraph("─" * 40)

        elif raw.startswith("- ") or raw.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            _add_inline_markup(para, raw[2:], _add_cn_font)

        elif re.match(r"^\[?\d+\]", raw) and in_ref_section:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.5)
            _add_inline_markup(para, raw, _add_cn_font)
            para.runs[-1].font.size = Pt(10) if para.runs else None

        elif raw.strip():
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进两字符
            _add_inline_markup(para, raw, _add_cn_font)

        else:
            doc.add_paragraph("")

    doc.save(str(path))
    logger.info(f"Word 导出成功：{path.resolve()}")
    return str(path.resolve())


def _add_inline_markup(para, text: str, set_font_fn):
    """处理行内 Markdown 标记（粗体、斜体、代码）并写入段落 run"""
    # 将 **bold**、*italic*、`code` 解析为 tokens
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
            set_font_fn(run)
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
            set_font_fn(run)
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
        elif part:
            run = para.add_run(part)
            set_font_fn(run)


# ============================================================
# 统一导出接口
# ============================================================

def export(content: str, fmt: str = "md", output_dir: str = "./output",
           filename: str = "survey") -> str:
    """
    统一导出接口，根据格式分发。

    Args:
        content   : Markdown 综述内容
        fmt       : "md" | "pdf" | "docx" | "word"
        output_dir: 输出目录
        filename  : 文件名（不含扩展名）

    Returns:
        str: 输出文件的绝对路径
    """
    fmt = fmt.lower().strip(".")
    os.makedirs(output_dir, exist_ok=True)

    ext_map = {"md": "md", "pdf": "pdf", "docx": "docx", "word": "docx"}
    if fmt not in ext_map:
        raise ValueError(f"不支持的导出格式：{fmt}。支持格式：md, pdf, docx")

    ext = ext_map[fmt]
    output_path = os.path.join(output_dir, f"{filename}.{ext}")

    dispatch = {
        "md":   export_markdown,
        "pdf":  export_pdf,
        "docx": export_docx,
        "word": export_docx,
    }
    return dispatch[fmt](content, output_path)


# ============================================================
# 快速测试
# ============================================================

if __name__ == "__main__":
    sample = """# 大型语言模型综述

## 1. 引言

大型语言模型（LLM）是近年来人工智能领域的重要突破，以 GPT、LLaMA 等为代表，推动了自然语言处理任务的全面提升 [1]。

## 2. 相关工作

**Transformer 架构** 由 Vaswani 等人于 2017 年提出 [2]，是现代 LLM 的基础结构。自注意力机制的引入显著提升了模型对长程依赖关系的建模能力。

### 2.1 预训练模型

*BERT* 采用双向编码器结构，在多个 NLP 基准上取得了突破性成果 [3]。

## 3. 结论

LLM 在多个基准任务上已超越人类平均水平，未来研究方向包括模型压缩、多模态融合和对齐技术。

---

## 参考文献

[1] Brown et al. *Language Models are Few-Shot Learners*. NeurIPS 2020.
[2] Vaswani et al. *Attention Is All You Need*. NeurIPS 2017.
[3] Devlin et al. *BERT: Pre-training of Deep Bidirectional Transformers*. NAACL 2019.
"""
    Path("./test_output").mkdir(exist_ok=True)

    p1 = export(sample, fmt="md",   output_dir="./test_output", filename="test_survey")
    print(f"✅ Markdown: {p1}")

    try:
        p2 = export(sample, fmt="pdf",  output_dir="./test_output", filename="test_survey")
        print(f"✅ PDF:      {p2}")
    except Exception as e:
        print(f"⚠️  PDF 跳过：{e}")

    try:
        p3 = export(sample, fmt="docx", output_dir="./test_output", filename="test_survey")
        print(f"✅ Word:     {p3}")
    except Exception as e:
        print(f"⚠️  Word 跳过：{e}")
